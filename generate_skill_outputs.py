"""Generate 3 output formats from a SKILL.md

Usage:
    python generate_skill_outputs.py login-paysage          # single skill
    python generate_skill_outputs.py --all                   # all skills

Reads SKILL.md → produces:
  - <name>.docx (Word)
  - <name>/obsidian/ (vault with tagged notes)
"""
import sys, re, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
SKILLS_DIR = ROOT / "open-design" / "skills"

# ── PARSER ──

def parse_skill(skill_path: Path) -> dict:
    """Extrait les sections d'un SKILL.md"""
    text = skill_path.read_text(encoding="utf-8")
    skill_name = skill_path.parent.name

    sections = {
        "name": skill_name,
        "fonction": "",
        "contraintes_globales": [],
        "sous_systemes": [],
        "exemples": [],
    }

    # Extraire fonction principale
    fn_match = re.search(r"##\s+1\.\s+Fonction Principale.*?(?=##\s+2\.|\Z)", text, re.DOTALL)
    if fn_match:
        sections["fonction"] = fn_match.group(0).strip()

    # Extraire type de système
    sections["systeme_type"] = "systeme-ferme"
    if "ouvert" in sections["fonction"].lower():
        sections["systeme_type"] = "systeme-ouvert"

    # Extraire contraintes globales
    cg_match = re.search(r"### Tableau global.*?(?=### Sous-système|\Z)", text, re.DOTALL)
    if cg_match:
        rows = re.findall(r"\|\s*(C\d+)\s*\|\s*(.+?)\s*\|", cg_match.group(0))
        sections["contraintes_globales"] = [{"id": r[0], "text": r[1]} for r in rows]

    # Extraire sous-systèmes
    for ss_match in re.finditer(r"### Sous-système (\w+) — (.+?)\n", text):
        ss_letter = ss_match.group(1)
        ss_name = ss_match.group(2)
        body = text[ss_match.end():]
        body = body[:body.find("### Sous-système")] if "### Sous-système" in body else body[:body.find("## ")] if "## " in body else body
        rows = re.findall(r"\|\s*([A-Z]\d+)\s*\|\s*(.+?)\s*\|", body)
        sections["sous_systemes"].append({
            "id": ss_letter,
            "name": ss_name,
            "constraints": [{"id": r[0], "text": r[1]} for r in rows]
        })

    # Extraire exemples
    ex_match = re.search(r"##\s+3\.\s+Deux exemples.*", text, re.DOTALL)
    if ex_match:
        sections["exemples_raw"] = ex_match.group(0).strip()

    return sections

# ── WORD GENERATOR ──

def generate_docx(sections: dict, output_path: Path):
    doc = Document()

    # Title
    title = doc.add_heading(f"Skill: {sections['name']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Type: {'Ouvert' if 'ouvert' in sections.get('systeme_type','') else 'Fermé'}")
    doc.add_paragraph(f"Tags: skill, {sections['name']}, {sections.get('systeme_type','')}, priorite-1")

    # 1. Fonction
    doc.add_heading("1. Fonction Principale", level=1)
    for line in sections["fonction"].split("\n"):
        if not line.startswith("#") and line.strip():
            doc.add_paragraph(line.strip())

    # 2. Contraintes
    doc.add_heading("2. Contraintes Fonctionnelles", level=1)

    if sections["contraintes_globales"]:
        doc.add_heading("Tableau global", level=2)
        table = doc.add_table(rows=1 + len(sections["contraintes_globales"]), cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "#"
        table.rows[0].cells[1].text = "Contrainte"
        for i, c in enumerate(sections["contraintes_globales"], 1):
            table.rows[i].cells[0].text = c["id"]
            table.rows[i].cells[1].text = c["text"]

    for ss in sections["sous_systemes"]:
        doc.add_heading(f"Sous-système {ss['id']} — {ss['name']}", level=2)
        table = doc.add_table(rows=1 + len(ss["constraints"]), cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "#"
        table.rows[0].cells[1].text = "Contrainte"
        for i, c in enumerate(ss["constraints"], 1):
            table.rows[i].cells[0].text = c["id"]
            table.rows[i].cells[1].text = c["text"]

    # 3. Exemples
    doc.add_heading("3. Exemples", level=1)
    if "exemples_raw" in sections:
        for line in sections["exemples_raw"].split("\n")[:50]:
            if line.strip():
                doc.add_paragraph(line.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

# ── OBSIDIAN GENERATOR ──

def generate_obsidian(sections: dict, output_dir: Path):
    output_dir.mkdir(parents=True, exist_ok=True)
    name = sections["name"]
    stype = sections.get("systeme_type", "systeme-ferme")

    def write_note(filename: str, content: str, tags: list):
        tags_yml = "\n".join(f"  - {t}" for t in tags)
        note = f"---\ntags:\n{tags_yml}\n---\n\n{content}"
        (output_dir / filename).write_text(note, encoding="utf-8")

    # index.md (hub)
    links = []
    for fname, label in [("1-fonction", "Fonction principale"),
                          ("2-contraintes", "Contraintes"),
                          ("3-exemples", "Exemples")]:
        links.append(f"- [[{fname}|{label}]]")
    hub = f"# {name} — Hub\n\n" + "\n".join(links)
    write_note("index.md", hub, ["skill", name, "hub"])

    # 1-fonction.md
    fn_content = f"# Fonction Principale\n\nType: **{stype}**\n\n"
    fn_content += sections.get("fonction", "").replace("### Type : Système Fermé", "").replace("### Type : Système Ouvert", "")
    write_note("1-fonction.md", fn_content, ["skill", name, "fonction", stype])

    # 2-contraintes.md
    cg_text = "# Contraintes Fonctionnelles\n\n"
    if sections["contraintes_globales"]:
        cg_text += "## Tableau global\n\n"
        for c in sections["contraintes_globales"]:
            cg_text += f"- **{c['id']}**: {c['text']}\n"
        cg_text += "\n"

    for ss in sections["sous_systemes"]:
        ss_file = f"2{ss['id'].lower()}-{ss['name'].lower().replace(' ','-').replace('/','-')[:30]}"
        cg_text += f"## [[{ss_file}|Sous-système {ss['id']}: {ss['name']}]]\n\n"
        # Create subsystem note
        ss_content = f"# Sous-système {ss['id']}: {ss['name']}\n\n"
        for c in ss["constraints"]:
            ss_content += f"- **{c['id']}**: {c['text']}\n"
        write_note(f"{ss_file}.md", ss_content, ["skill", name, "contrainte", "sous-systeme", f"priorite-{min(int(ss['id'].encode()[0])%3+1,3)}"])

    write_note("2-contraintes.md", cg_text, ["skill", name, "contrainte"])

    # 3-exemples.md
    ex_text = "# Exemples\n\n" + sections.get("exemples_raw", "Aucun exemple.").replace("## 3. Deux exemples\n\n", "")
    write_note("3-exemples.md", ex_text[:5000], ["skill", name, "exemple"])

# ── MAIN ──

def process_skill(skill_dir: Path):
    skill_md = skill_dir / "SKILL.md"
    if not skill_md.exists():
        print(f"  SKIP {skill_dir.name}: pas de SKILL.md")
        return

    sections = parse_skill(skill_md)
    name = sections["name"]

    # DOCX
    docx_path = skill_dir / f"{name}.docx"
    generate_docx(sections, docx_path)
    print(f"  DOCX: {docx_path.name}")

    # OBSIDIAN
    obs_dir = skill_dir / "obsidian"
    generate_obsidian(sections, obs_dir)
    note_count = len(list(obs_dir.glob("*.md")))
    print(f"  OBSIDIAN: {note_count} notes dans {obs_dir.relative_to(ROOT)}")

def main():
    args = sys.argv[1:]

    if not args:
        print("Usage: python generate_skill_outputs.py <nom-du-skill> | --all")
        return

    if "--all" in args:
        for skill_dir in sorted(SKILLS_DIR.iterdir()):
            if skill_dir.is_dir() and (skill_dir / "SKILL.md").exists():
                print(f"\n{skill_dir.name}:")
                process_skill(skill_dir)
        print("\nDone - all skills processed")
    else:
        skill_dir = SKILLS_DIR / args[0]
        if skill_dir.exists():
            process_skill(skill_dir)
        else:
            print(f"Skill '{args[0]}' not found in {SKILLS_DIR}")

if __name__ == "__main__":
    main()
